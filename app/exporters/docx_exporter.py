"""DOCX report export."""

from __future__ import annotations

from app.config import settings
from app.core.exceptions import FileProcessingError
from app.utils.file_utils import ensure_parent, make_job_id, safe_filename


def export_docx(title: str, sections: list[tuple[str, str]], name: str, job_id: str | None = None) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise FileProcessingError("DOCX rapor için python-docx paketi gerekli.") from exc

    job = job_id or make_job_id("report")
    output = settings.REPORT_DIR / job / f"{safe_filename(name)}.docx"
    ensure_parent(output)

    doc = Document()
    doc.add_heading(title, level=1)
    for heading, body in sections:
        doc.add_heading(heading, level=2)
        for paragraph in body.splitlines():
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    doc.save(output)
    return str(output.resolve())
