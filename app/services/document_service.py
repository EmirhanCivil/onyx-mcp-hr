"""Generic document builder — any format from a single tool call.

Supported formats:
  xlsx | csv | tsv | txt | md | html | json | docx | pdf

Each format accepts a small, format-specific content schema (described in the
docstring of `create`). All outputs land under settings.OUTPUT_DIR/documents/<job>/
so the existing nginx sidecar at :8007 serves them and `_to_public_url()` in
response_schema.py rewrites the path to a markdown-ready URL automatically.
"""

from __future__ import annotations

import csv
import io
import json
from html import escape
from pathlib import Path
from typing import Any, Iterable, Sequence

from app.config import settings
from app.core.exceptions import AppError
from app.utils.file_utils import ensure_parent, make_job_id, safe_filename


_DOC_ROOT = settings.OUTPUT_DIR / "documents"

_SUPPORTED = {"xlsx", "csv", "tsv", "txt", "md", "html", "json", "docx", "pdf"}

_PDF_PAGE_SIZES = {"A4", "LETTER", "LEGAL", "A3", "A5"}


def _job_dir() -> Path:
    job = make_job_id("doc")
    out = _DOC_ROOT / job
    out.mkdir(parents=True, exist_ok=True)
    return out


def _resolve_output(filename: str, fmt: str) -> Path:
    base = safe_filename(filename or f"document.{fmt}")
    if not base.lower().endswith(f".{fmt}"):
        base = f"{base}.{fmt}"
    target = _job_dir() / base
    ensure_parent(target)
    return target


_XLSX_INVALID_SHEET_CHARS = set("\\/?*[]:")


def _xlsx_sheet_name(value: str, fallback: str) -> str:
    raw = (str(value or fallback) or fallback).strip()
    cleaned = "".join(c for c in raw if c not in _XLSX_INVALID_SHEET_CHARS)
    cleaned = cleaned.strip()[:31] or fallback
    return cleaned


def _is_numeric(value: Any) -> bool:
    if isinstance(value, bool):
        return False
    return isinstance(value, (int, float))


_PDF_FONTS_REGISTERED: dict[str, str] | None = None


def _ensure_unicode_fonts() -> tuple[str, str, str]:
    """Register DejaVu Sans (Unicode-capable) for ReportLab. Returns (regular, bold, italic) names."""
    global _PDF_FONTS_REGISTERED
    if _PDF_FONTS_REGISTERED:
        f = _PDF_FONTS_REGISTERED
        return f["regular"], f["bold"], f["italic"]

    candidate_dirs = [
        Path("/usr/local/lib/python3.11/site-packages/matplotlib/mpl-data/fonts/ttf"),
        Path("/usr/share/fonts/truetype/dejavu"),
        Path("/usr/share/fonts/truetype"),
        Path("/Library/Fonts"),
        Path("C:/Windows/Fonts"),
    ]
    files = {"regular": "DejaVuSans.ttf", "bold": "DejaVuSans-Bold.ttf",
             "italic": "DejaVuSans-Oblique.ttf", "bolditalic": "DejaVuSans-BoldOblique.ttf"}
    found: dict[str, Path] = {}
    for key, fname in files.items():
        for d in candidate_dirs:
            p = d / fname
            if p.exists():
                found[key] = p
                break

    if "regular" not in found:
        # Fallback — keep Helvetica (Latin-1 only)
        _PDF_FONTS_REGISTERED = {"regular": "Helvetica", "bold": "Helvetica-Bold", "italic": "Helvetica-Oblique"}
        return "Helvetica", "Helvetica-Bold", "Helvetica-Oblique"

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.fonts import addMapping

    family = "AppSans"
    regular = f"{family}"
    bold = f"{family}-Bold"
    italic = f"{family}-Italic"
    bolditalic = f"{family}-BoldItalic"

    pdfmetrics.registerFont(TTFont(regular, str(found["regular"])))
    if "bold" in found:
        pdfmetrics.registerFont(TTFont(bold, str(found["bold"])))
    else:
        bold = regular
    if "italic" in found:
        pdfmetrics.registerFont(TTFont(italic, str(found["italic"])))
    else:
        italic = regular
    if "bolditalic" in found:
        pdfmetrics.registerFont(TTFont(bolditalic, str(found["bolditalic"])))
    else:
        bolditalic = bold

    addMapping(family, 0, 0, regular)
    addMapping(family, 1, 0, bold)
    addMapping(family, 0, 1, italic)
    addMapping(family, 1, 1, bolditalic)

    _PDF_FONTS_REGISTERED = {"regular": regular, "bold": bold, "italic": italic}
    return regular, bold, italic


def _looks_numeric(value: Any) -> bool:
    if _is_numeric(value):
        return True
    if value is None or value == "":
        return False
    if isinstance(value, str):
        s = value.strip().replace(" ", "").replace(",", ".").replace("%", "")
        if s.startswith(("+", "-")):
            s = s[1:]
        try:
            float(s)
            return True
        except ValueError:
            return False
    return False


def _coerce_rows(rows: Any) -> list[list[Any]]:
    if rows is None:
        return []
    if not isinstance(rows, list):
        raise AppError("BAD_CONTENT", "rows must be a list of lists.")
    coerced: list[list[Any]] = []
    for row in rows:
        if isinstance(row, dict):
            coerced.append(list(row.values()))
        elif isinstance(row, (list, tuple)):
            coerced.append(list(row))
        else:
            coerced.append([row])
    return coerced


def _normalize_sheets(content: dict[str, Any]) -> list[dict[str, Any]]:
    if "sheets" in content and isinstance(content["sheets"], list):
        return content["sheets"]
    rows = content.get("rows")
    if rows is not None:
        return [
            {
                "name": content.get("sheet_name") or content.get("name") or "Sayfa1",
                "columns": content.get("columns") or [],
                "rows": rows,
            }
        ]
    return []


# ---------------------------------------------------------------------------
# Format builders
# ---------------------------------------------------------------------------


def _build_xlsx(target: Path, content: dict[str, Any], doc_title: str = "") -> dict[str, Any]:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    sheets = _normalize_sheets(content)
    if not sheets:
        raise AppError(
            "BAD_CONTENT",
            "xlsx için content.sheets veya content.rows alanı gerekli.",
        )

    wb = Workbook()
    wb.remove(wb.active)

    title_font = Font(bold=True, color="FFFFFF", name="Calibri", size=15)
    title_fill = PatternFill("solid", fgColor="134E4A")
    title_align = Alignment(horizontal="center", vertical="center")
    header_font = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
    header_fill = PatternFill("solid", fgColor="0F766E")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    text_align = Alignment(horizontal="left", vertical="center", wrap_text=True, indent=1)
    num_align = Alignment(horizontal="right", vertical="center", indent=1)
    int_align = Alignment(horizontal="center", vertical="center")
    first_col_font = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
    first_col_fill = PatternFill("solid", fgColor="0F766E")
    zebra_fill = PatternFill("solid", fgColor="F1F5F9")
    grid_color = "94A3B8"
    outer_color = "0F766E"
    grid_side = Side(style="thin", color=grid_color)
    outer_side = Side(style="medium", color=outer_color)
    cell_border = Border(left=grid_side, right=grid_side, top=grid_side, bottom=grid_side)

    sheet_meta: list[dict[str, Any]] = []
    for idx, sheet in enumerate(sheets):
        if not isinstance(sheet, dict):
            raise AppError("BAD_CONTENT", "Her sheet bir nesne olmalı.")
        fallback = f"Sayfa{idx + 1}"
        name = _xlsx_sheet_name(sheet.get("name") or fallback, fallback)
        ws = wb.create_sheet(title=name)
        columns = sheet.get("columns") or []
        rows = _coerce_rows(sheet.get("rows"))
        sheet_title = sheet.get("title") or (doc_title if idx == 0 else "")
        first_col_emphasis = bool(sheet.get("first_column_emphasis", True))

        col_count = max(len(columns), max((len(r) for r in rows), default=0))
        if col_count == 0:
            sheet_meta.append({"name": name, "columns": [], "rows": 0})
            continue

        cursor = 1
        table_top = 1
        title_row = 0

        if sheet_title:
            title_row = cursor
            ws.merge_cells(start_row=cursor, start_column=1, end_row=cursor, end_column=col_count)
            cell = ws.cell(row=cursor, column=1, value=str(sheet_title))
            cell.font = title_font
            cell.fill = title_fill
            cell.alignment = title_align
            ws.row_dimensions[cursor].height = 36
            cursor += 1
            table_top = cursor
            ws.row_dimensions[cursor].height = 6  # spacing
            cursor += 1

        header_row = 0
        if columns:
            header_row = cursor
            table_top = header_row if not sheet_title else table_top
            for col_idx, header in enumerate(columns, start=1):
                cell = ws.cell(row=cursor, column=col_idx, value=str(header))
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align
                cell.border = cell_border
            ws.row_dimensions[cursor].height = 30
            cursor += 1
            ws.freeze_panes = ws.cell(row=cursor, column=1).coordinate

        data_start = cursor
        col_numeric_flag = [True] * col_count  # all-numeric column?
        col_has_value = [False] * col_count
        max_lens = [len(str(c)) for c in columns] + [4] * (col_count - len(columns))

        # First pass: detect numeric-only columns
        for row in rows:
            for c_idx in range(col_count):
                value = row[c_idx] if c_idx < len(row) else None
                if value is None or value == "":
                    continue
                col_has_value[c_idx] = True
                if not _is_numeric(value):
                    col_numeric_flag[c_idx] = False

        for r_offset, row in enumerate(rows):
            r_idx = data_start + r_offset
            zebra = r_offset % 2 == 1
            for c_idx in range(1, col_count + 1):
                value = row[c_idx - 1] if c_idx - 1 < len(row) else None
                cell = ws.cell(row=r_idx, column=c_idx, value=value)
                cell.border = cell_border
                if _is_numeric(value):
                    if isinstance(value, float) and not value.is_integer():
                        cell.number_format = "#,##0.00"
                        cell.alignment = num_align
                    else:
                        cell.number_format = "#,##0" if abs(float(value)) >= 1000 else "0"
                        cell.alignment = int_align
                else:
                    cell.alignment = text_align
                if first_col_emphasis and c_idx == 1 and columns:
                    cell.font = first_col_font
                    cell.fill = first_col_fill
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                elif zebra:
                    cell.fill = zebra_fill
                if value is not None:
                    max_lens[c_idx - 1] = max(max_lens[c_idx - 1], len(str(value)))
            ws.row_dimensions[r_idx].height = 22

        last_data_row = data_start + len(rows) - 1 if rows else data_start - 1

        # Outer thick border around entire table (header + data)
        if columns and rows:
            top_row = header_row
            for c_idx in range(1, col_count + 1):
                # top edge
                top_cell = ws.cell(row=top_row, column=c_idx)
                tb = top_cell.border
                top_cell.border = Border(left=tb.left, right=tb.right, top=outer_side, bottom=tb.bottom)
                # bottom edge
                bot_cell = ws.cell(row=last_data_row, column=c_idx)
                bb = bot_cell.border
                bot_cell.border = Border(left=bb.left, right=bb.right, top=bb.top, bottom=outer_side)
            for r_idx in range(top_row, last_data_row + 1):
                # left edge
                lc = ws.cell(row=r_idx, column=1)
                lb = lc.border
                lc.border = Border(left=outer_side, right=lb.right, top=lb.top, bottom=lb.bottom)
                # right edge
                rc = ws.cell(row=r_idx, column=col_count)
                rb = rc.border
                rc.border = Border(left=rb.left, right=outer_side, top=rb.top, bottom=rb.bottom)
            # Re-apply corners
            corners = [(top_row, 1, "tl"), (top_row, col_count, "tr"),
                       (last_data_row, 1, "bl"), (last_data_row, col_count, "br")]
            for r_idx, c_idx, pos in corners:
                cell = ws.cell(row=r_idx, column=c_idx)
                bb = cell.border
                cell.border = Border(
                    left=outer_side if pos in {"tl", "bl"} else bb.left,
                    right=outer_side if pos in {"tr", "br"} else bb.right,
                    top=outer_side if pos in {"tl", "tr"} else bb.top,
                    bottom=outer_side if pos in {"bl", "br"} else bb.bottom,
                )

        for c_idx in range(1, col_count + 1):
            letter = get_column_letter(c_idx)
            header_len = len(str(columns[c_idx - 1])) if c_idx - 1 < len(columns) else 0
            data_len = max_lens[c_idx - 1]
            target_width = max(12, min(48, max(header_len, data_len) + 6))
            ws.column_dimensions[letter].width = target_width

        if header_row and rows:
            ws.auto_filter.ref = f"{ws.cell(row=header_row, column=1).coordinate}:{ws.cell(row=last_data_row, column=col_count).coordinate}"

        sheet_meta.append({"name": name, "columns": list(columns), "rows": len(rows)})

    wb.save(target)
    return {"sheets": sheet_meta}


def _build_csv(target: Path, content: dict[str, Any], delimiter: str) -> dict[str, Any]:
    columns = content.get("columns") or []
    rows = _coerce_rows(content.get("rows"))
    if not columns and not rows:
        raise AppError("BAD_CONTENT", "csv için columns ve/veya rows gerekli.")
    custom = content.get("delimiter")
    if isinstance(custom, str) and custom:
        delimiter = "\t" if custom.lower() in {"tab", "\\t"} else custom[0]
    # NOTE: UTF-8 BOM + ; delimiter (TR Excel native). `sep=` directive AVOIDED —
    # it breaks BOM detection in Excel and causes Turkish characters to mojibake.
    with target.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.writer(fh, delimiter=delimiter, quoting=csv.QUOTE_MINIMAL, lineterminator="\r\n")
        if columns:
            writer.writerow([str(c) for c in columns])
        for row in rows:
            writer.writerow(["" if v is None else v for v in row])
    return {"columns": list(columns), "rows": len(rows), "delimiter": delimiter}


def _build_txt(target: Path, content: dict[str, Any]) -> dict[str, Any]:
    text = content.get("text")
    if text is None and "lines" in content:
        text = "\n".join(str(line) for line in content.get("lines") or [])
    if text is None:
        raise AppError("BAD_CONTENT", "txt için content.text veya content.lines gerekli.")
    target.write_text(str(text), encoding="utf-8")
    return {"chars": len(text)}


def _build_json(target: Path, content: dict[str, Any]) -> dict[str, Any]:
    payload = content.get("data", content)
    body = json.dumps(payload, ensure_ascii=False, indent=2, default=str)
    target.write_text(body, encoding="utf-8")
    return {"bytes": len(body.encode("utf-8"))}


def _md_table(columns: Sequence[Any], rows: Iterable[Sequence[Any]]) -> str:
    cols = [str(c) for c in columns]
    if not cols:
        return ""
    out = ["| " + " | ".join(cols) + " |", "| " + " | ".join(["---"] * len(cols)) + " |"]
    for row in rows:
        cells = [str(v) if v is not None else "" for v in row]
        out.append("| " + " | ".join(cells) + " |")
    return "\n".join(out)


def _build_md(target: Path, content: dict[str, Any], title: str) -> dict[str, Any]:
    parts: list[str] = []
    if title:
        parts.append(f"# {title}\n")

    if "markdown" in content and content["markdown"]:
        parts.append(str(content["markdown"]))
    else:
        for section in content.get("sections") or []:
            if not isinstance(section, dict):
                continue
            level = max(1, min(6, int(section.get("level", 2))))
            heading = section.get("heading")
            if heading:
                parts.append(f"\n{'#' * level} {heading}\n")
            body = section.get("body") or section.get("text")
            if body:
                parts.append(str(body))
            if section.get("table"):
                tbl = section["table"]
                parts.append("\n" + _md_table(tbl.get("columns") or [], _coerce_rows(tbl.get("rows"))))
            for bullet in section.get("bullets") or []:
                parts.append(f"- {bullet}")
        if content.get("table"):
            tbl = content["table"]
            parts.append(_md_table(tbl.get("columns") or [], _coerce_rows(tbl.get("rows"))))

    body = "\n\n".join(p for p in parts if p).strip() + "\n"
    target.write_text(body, encoding="utf-8")
    return {"chars": len(body)}


def _html_escape(value: Any) -> str:
    return escape("" if value is None else str(value))


def _build_html(target: Path, content: dict[str, Any], title: str) -> dict[str, Any]:
    body_parts: list[str] = []

    if "html" in content and content["html"]:
        body_parts.append(str(content["html"]))
    elif "markdown" in content and content["markdown"]:
        try:
            from markdown_it import MarkdownIt

            md = MarkdownIt("commonmark", {"html": False, "linkify": True, "typographer": True})
            md.enable("table")
            body_parts.append(md.render(str(content["markdown"])))
        except Exception:
            body_parts.append(f"<pre>{_html_escape(content['markdown'])}</pre>")
    else:
        for section in content.get("sections") or []:
            if not isinstance(section, dict):
                continue
            level = max(1, min(6, int(section.get("level", 2))))
            heading = section.get("heading")
            if heading:
                body_parts.append(f"<h{level}>{_html_escape(heading)}</h{level}>")
            body = section.get("body") or section.get("text")
            if body:
                for para in str(body).splitlines():
                    p = para.strip()
                    if p:
                        body_parts.append(f"<p>{_html_escape(p)}</p>")
            if section.get("bullets"):
                items = "".join(f"<li>{_html_escape(b)}</li>" for b in section["bullets"])
                body_parts.append(f"<ul>{items}</ul>")
            if section.get("table"):
                tbl = section["table"]
                body_parts.append(_html_table(tbl.get("columns") or [], _coerce_rows(tbl.get("rows"))))
        if content.get("table"):
            tbl = content["table"]
            body_parts.append(_html_table(tbl.get("columns") or [], _coerce_rows(tbl.get("rows"))))

    inner = "\n".join(body_parts) or f"<p>{_html_escape(content.get('text', ''))}</p>"
    page = _wrap_html(title or "Document", inner)
    target.write_text(page, encoding="utf-8")
    return {"chars": len(page)}


def _html_table(columns: Sequence[Any], rows: Iterable[Sequence[Any]]) -> str:
    head = "".join(f"<th>{_html_escape(c)}</th>" for c in columns)
    body = "".join(
        "<tr>" + "".join(f"<td>{_html_escape(v)}</td>" for v in row) + "</tr>"
        for row in rows
    )
    return f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>"


def _wrap_html(title: str, inner: str) -> str:
    return f"""<!doctype html>
<html lang="tr"><head>
<meta charset="utf-8" />
<meta name="viewport" content="width=device-width,initial-scale=1" />
<title>{_html_escape(title)}</title>
<style>
:root {{ color-scheme: light dark; }}
body {{ font-family: 'Inter', system-ui, -apple-system, 'Segoe UI', Roboto, sans-serif;
        max-width: 980px; margin: 0 auto; padding: 32px 22px; color: #0F172A;
        background: #F8FAFC; line-height: 1.6; }}
h1, h2, h3, h4 {{ color: #134E4A; }}
table {{ border-collapse: collapse; width: 100%; margin: 14px 0; }}
th, td {{ border: 1px solid #E2E8F0; padding: 8px 10px; text-align: left; }}
th {{ background: #0F766E; color: #fff; }}
tr:nth-child(even) td {{ background: #F1F5F9; }}
code, pre {{ background: #F1F5F9; padding: 2px 6px; border-radius: 4px; }}
pre {{ padding: 12px; overflow:auto; }}
@media (prefers-color-scheme: dark) {{
  body {{ background:#0B1120; color:#E2E8F0; }}
  h1,h2,h3,h4 {{ color:#5EEAD4; }}
  th {{ background:#0F766E; }} td {{ border-color:#1E293B; }}
  tr:nth-child(even) td {{ background:#111827; }}
  code, pre {{ background:#111827; }}
}}
</style></head>
<body>
<h1>{_html_escape(title)}</h1>
{inner}
</body></html>"""


def _build_docx(target: Path, content: dict[str, Any], title: str) -> dict[str, Any]:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise AppError("MISSING_DEP", "DOCX için python-docx paketi gerekli.") from exc

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)
        try:
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x13, 0x4E, 0x4A)
        except Exception:
            pass

    sections = content.get("sections") or []
    if not sections and (content.get("text") or content.get("paragraphs")):
        text = content.get("text") or "\n".join(content.get("paragraphs") or [])
        sections = [{"body": text}]
    if not sections and content.get("table"):
        sections = [{"table": content["table"]}]

    for section in sections:
        if not isinstance(section, dict):
            continue
        if section.get("heading"):
            level = max(1, min(4, int(section.get("level", 1))))
            doc.add_heading(str(section["heading"]), level=level)
        body = section.get("body") or section.get("text")
        if body:
            for para in str(body).splitlines():
                p = para.strip()
                if p:
                    doc.add_paragraph(p)
        for bullet in section.get("bullets") or []:
            doc.add_paragraph(str(bullet), style="List Bullet")
        for number in section.get("numbered") or []:
            doc.add_paragraph(str(number), style="List Number")
        if section.get("table"):
            _docx_add_styled_table(doc, section["table"])

    doc.save(target)
    return {"sections": len(sections)}


def _docx_add_styled_table(doc: Any, tbl: dict[str, Any]) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    cols = list(tbl.get("columns") or [])
    rows = _coerce_rows(tbl.get("rows"))
    col_count = len(cols) or (max((len(r) for r in rows), default=0))
    if not col_count:
        return

    has_header = bool(cols)
    table = doc.add_table(rows=(1 if has_header else 0) + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    def _shade(cell: Any, color_hex: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex)
        tc_pr.append(shd)

    if has_header:
        hdr = table.rows[0].cells
        for i, c in enumerate(cols):
            hdr[i].text = str(c)
            _shade(hdr[i], "0F766E")
            for paragraph in hdr[i].paragraphs:
                paragraph.alignment = 1  # center
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10.5)

    for r_idx, row in enumerate(rows):
        target_row = table.rows[(1 if has_header else 0) + r_idx]
        cells = target_row.cells
        for c_idx in range(col_count):
            value = row[c_idx] if c_idx < len(row) else None
            cells[c_idx].text = "" if value is None else str(value)
            if r_idx % 2 == 1:
                _shade(cells[c_idx], "F1F5F9")
            for paragraph in cells[c_idx].paragraphs:
                if _looks_numeric(value):
                    paragraph.alignment = 1  # center
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def _build_pdf(target: Path, content: dict[str, Any], title: str, page_size: str) -> dict[str, Any]:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A3, A4, A5, LEGAL, LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:
        raise AppError("MISSING_DEP", "PDF için reportlab paketi gerekli.") from exc

    base_font, bold_font, italic_font = _ensure_unicode_fonts()

    size_map = {"A4": A4, "A3": A3, "A5": A5, "LETTER": LETTER, "LEGAL": LEGAL}
    page = size_map.get((page_size or "A4").upper(), A4)

    doc = SimpleDocTemplate(
        str(target),
        pagesize=page,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=title or "Document",
    )

    base = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle",
        parent=base["Title"],
        fontName=bold_font,
        textColor=colors.HexColor("#134E4A"),
        fontSize=22,
        leading=26,
        spaceAfter=14,
    )
    h2 = ParagraphStyle(
        "DocH2",
        parent=base["Heading2"],
        fontName=bold_font,
        textColor=colors.HexColor("#0F766E"),
        spaceBefore=10,
        spaceAfter=4,
    )
    body = ParagraphStyle(
        "DocBody",
        parent=base["BodyText"],
        fontName=base_font,
        fontSize=10.5,
        leading=15,
        spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        "DocBullet", parent=body, fontName=base_font, leftIndent=14, bulletIndent=4
    )

    flow: list[Any] = []
    if title:
        flow.append(Paragraph(_html_escape(title), title_style))

    sections = content.get("sections") or []
    if not sections and (content.get("text") or content.get("paragraphs")):
        text = content.get("text") or "\n".join(content.get("paragraphs") or [])
        sections = [{"body": text}]
    if not sections and content.get("table"):
        sections = [{"table": content["table"]}]

    for section in sections:
        if not isinstance(section, dict):
            continue
        if section.get("heading"):
            flow.append(Paragraph(_html_escape(section["heading"]), h2))
        section_body = section.get("body") or section.get("text")
        if section_body:
            for para in str(section_body).splitlines():
                p = para.strip()
                if p:
                    flow.append(Paragraph(_html_escape(p), body))
        for bullet in section.get("bullets") or []:
            flow.append(Paragraph(f"• {_html_escape(bullet)}", bullet_style))
        if section.get("table"):
            tbl = section["table"]
            cols = list(tbl.get("columns") or [])
            rows = _coerce_rows(tbl.get("rows"))
            if cols or rows:
                table_data = []
                if cols:
                    table_data.append([str(c) for c in cols])
                for row in rows:
                    table_data.append(["" if v is None else str(v) for v in row])
                t = Table(table_data, repeatRows=1 if cols else 0, hAlign="LEFT")
                style_cmds: list[Any] = [
                    ("FONTNAME", (0, 0), (-1, -1), base_font),
                    ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 7),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94A3B8")),
                    ("BOX", (0, 0), (-1, -1), 1.2, colors.HexColor("#0F766E")),
                ]
                if cols:
                    style_cmds.extend([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F766E")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), bold_font),
                        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                        ("LINEBELOW", (0, 0), (-1, 0), 1.0, colors.HexColor("#134E4A")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
                    ])
                else:
                    style_cmds.append(("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]))
                # Right-align numeric columns
                for c_idx in range(len(table_data[0]) if table_data else 0):
                    col_values = [r[c_idx] for r in table_data[1:]] if cols else [r[c_idx] for r in table_data]
                    if col_values and all(_looks_numeric(v) for v in col_values):
                        first_data_row = 1 if cols else 0
                        style_cmds.append(("ALIGN", (c_idx, first_data_row), (c_idx, -1), "CENTER"))
                t.setStyle(TableStyle(style_cmds))
                flow.append(Spacer(1, 6))
                flow.append(t)
                flow.append(Spacer(1, 6))

    if not flow:
        flow.append(Paragraph(_html_escape(content.get("text", "")), body))

    doc.build(flow)
    return {"page_size": page_size or "A4", "sections": len(sections)}


# ---------------------------------------------------------------------------
# Public service
# ---------------------------------------------------------------------------


class DocumentService:
    """Single-entry generic document builder."""

    SUPPORTED_FORMATS: set[str] = _SUPPORTED

    def create(
        self,
        format: str,
        filename: str = "",
        title: str = "",
        content: dict[str, Any] | None = None,
        page_size: str = "A4",
    ) -> dict[str, Any]:
        """Create a document file.

        Args:
            format: One of xlsx | csv | tsv | txt | md | html | json | docx | pdf
            filename: Output filename (extension auto-appended if missing).
            title: Optional document title (used by docx/pdf/html/md headings).
            content: Format-specific payload. Schemas:
                - xlsx: {sheets:[{name,columns,rows}]} OR {columns,rows,sheet_name?}
                - csv/tsv: {columns:[...], rows:[[...]]}
                - txt: {text:"..."} OR {lines:[...]}
                - json: {data: <any>} OR raw dict (used as-is)
                - md: {markdown:"..."} OR {sections:[{heading,level,body,bullets,table}]}
                - html: {html:"..."} OR {markdown:"..."} OR sections schema
                - docx: {sections:[{heading,level,body,bullets,numbered,table}]}
                        OR {text:"..."} / {paragraphs:[...]}
                - pdf: same as docx schema
            page_size: PDF page size — A4|LETTER|LEGAL|A3|A5 (default A4).

        Returns:
            dict with keys: path, format, filename, title, meta
        """
        fmt = (format or "").lower().strip()
        if fmt == "yaml":
            fmt = "yml"
        if fmt not in _SUPPORTED:
            raise AppError(
                "UNSUPPORTED_FORMAT",
                f"format='{format}' desteklenmiyor. Desteklenenler: {sorted(_SUPPORTED)}",
            )

        body = content if isinstance(content, dict) else {}
        if not body and fmt not in {"txt", "md", "html"}:
            raise AppError("BAD_CONTENT", "content alanı boş; format için en az bir alan gerekli.")

        target = _resolve_output(filename or (title or f"document.{fmt}"), fmt)
        title_clean = (title or "").strip()

        if fmt == "xlsx":
            meta = _build_xlsx(target, body, title_clean)
        elif fmt == "csv":
            meta = _build_csv(target, body, ";")
        elif fmt == "tsv":
            meta = _build_csv(target, body, "\t")
        elif fmt == "txt":
            meta = _build_txt(target, body or {"text": title_clean})
        elif fmt == "json":
            meta = _build_json(target, body)
        elif fmt == "md":
            meta = _build_md(target, body, title_clean)
        elif fmt == "html":
            meta = _build_html(target, body, title_clean)
        elif fmt == "docx":
            meta = _build_docx(target, body, title_clean)
        elif fmt == "pdf":
            meta = _build_pdf(target, body, title_clean, page_size)
        else:  # pragma: no cover
            raise AppError("UNSUPPORTED_FORMAT", f"format='{fmt}'")

        return {
            "path": str(target.resolve()),
            "format": fmt,
            "filename": target.name,
            "title": title_clean,
            "meta": meta,
        }


document_service = DocumentService()
